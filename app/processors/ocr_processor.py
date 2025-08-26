import easyocr
import fitz  # PyMuPDF
from PIL import Image
import io
import os
import logging
from typing import Dict, Any
import time
from pathlib import Path
import numpy as np

# Fix for Pillow 10.0.0+ compatibility with EasyOCR
if not hasattr(Image, 'ANTIALIAS'):
    Image.ANTIALIAS = Image.LANCZOS

from app.config import settings

logger = logging.getLogger(__name__)

class OCRProcessor:
    def __init__(self):
        """Initialize OCR with configured languages"""
        try:
            self.reader = easyocr.Reader(
                settings.ocr_languages, 
                gpu=settings.ocr_gpu
            )
            logger.info(f"OCR processor initialized with languages: {settings.ocr_languages}")
        except Exception as e:
            logger.error(f"Failed to initialize OCR processor: {e}")
            raise e
    
    async def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Main text extraction method"""
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        logger.info(f"Extracting text from {file_path} (type: {file_extension})")
        
        try:
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png']:
                return await self._extract_from_image(file_path)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_from_word(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise e
    
    async def _extract_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF and OCR"""
        
        start_time = time.time()
        doc = fitz.open(pdf_path)
        text_content = ""
        ocr_content = ""
        total_confidence = 0
        confidence_count = 0
        
        try:
            for page_num in range(doc.page_count):
                page = doc[page_num]
                
                # Extract direct text first
                page_text = page.get_text()
                if page_text.strip():
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                # If no text or minimal text, use OCR
                if len(page_text.strip()) < 50:
                    logger.debug(f"Using OCR for page {page_num + 1}")
                    
                    # Get page as image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom
                    img_data = pix.tobytes("png")
                    
                    # Perform OCR using image bytes directly
                    ocr_results = self.reader.readtext(img_data)
                    
                    page_ocr_text = ""
                    for (bbox, text, confidence) in ocr_results:
                        if confidence > 0.5:
                            page_ocr_text += f" {text}"
                            total_confidence += confidence
                            confidence_count += 1
                    
                    if page_ocr_text.strip():
                        ocr_content += f"\n--- Page {page_num + 1} (OCR) ---\n{page_ocr_text}"
            
            # Combine text
            combined_text = text_content + ocr_content
            avg_confidence = total_confidence / confidence_count if confidence_count > 0 else 0.0
            processing_time = time.time() - start_time
            
            result = {
                'text': combined_text.strip(),
                'confidence': avg_confidence,
                'pages_count': doc.page_count,
                'processing_time': processing_time,
                'language': self._detect_language(combined_text),
                'method': 'pdf_hybrid'
            }
            
            logger.info(f"PDF processed in {processing_time:.2f}s, {len(combined_text)} chars extracted")
            return result
            
        finally:
            doc.close()
    
    async def _extract_from_image(self, image_path: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        
        start_time = time.time()
        
        try:
            # Use file path directly for EasyOCR - more efficient and avoids type issues
            results = self.reader.readtext(image_path)
            
            extracted_text = ""
            confidence_scores = []
            
            for (bbox, text, confidence) in results:
                if confidence > 0.4:
                    extracted_text += f" {text}"
                    confidence_scores.append(confidence)
            
            avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0
            processing_time = time.time() - start_time
            
            result = {
                'text': extracted_text.strip(),
                'confidence': avg_confidence,
                'pages_count': 1,
                'processing_time': processing_time,
                'language': self._detect_language(extracted_text),
                'method': 'image_ocr',
                'detections_count': len(results),
                'high_confidence_count': len(confidence_scores)
            }
            
            logger.info(f"Image processed in {processing_time:.2f}s, {len(extracted_text)} chars extracted")
            return result
            
        except Exception as e:
            logger.error(f"Error processing image {image_path}: {e}")
            raise e
    
    async def _extract_from_word(self, word_path: str) -> Dict[str, Any]:
        """Extract text from Word document"""
        
        start_time = time.time()
        
        try:
            from docx import Document
            
            doc = Document(word_path)
            text_content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
            
            processing_time = time.time() - start_time
            
            result = {
                'text': text_content.strip(),
                'confidence': 1.0,  # Direct text extraction
                'pages_count': 1,
                'processing_time': processing_time,
                'language': self._detect_language(text_content),
                'method': 'word_direct'
            }
            
            logger.info(f"Word document processed in {processing_time:.2f}s, {len(text_content)} chars extracted")
            return result
            
        except Exception as e:
            logger.error(f"Error processing Word document {word_path}: {e}")
            raise e
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection"""
        if not text:
            return 'unknown'
        
        # Count Vietnamese characters
        vietnamese_chars = sum(1 for c in text if ord(c) > 127)
        total_chars = len(text.replace(' ', ''))
        
        if total_chars == 0:
            return 'unknown'
        
        vietnamese_ratio = vietnamese_chars / total_chars
        
        if vietnamese_ratio > 0.1:
            return 'vietnamese'
        else:
            return 'english'